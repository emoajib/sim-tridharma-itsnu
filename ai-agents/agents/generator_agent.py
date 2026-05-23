import json
import logging
import os
from datetime import datetime
import google.generativeai as genai

from sqlalchemy import text
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.base_agent import BaseAgent
from database import SessionLocal
from config import GENAI_API_KEY, GENAI_MODEL

logger = logging.getLogger("agent.generator")


class GeneratorAgent(BaseAgent):
    name = "generator"
    version = "1.1.0"  # Updated version for Gemini LLM integration

    def execute(self, data: dict) -> dict:
        missing = self.validate_input(data, ["prodi_id", "periode_id", "jenis_dokumen"])
        if missing:
            return {"status": "error", "message": f"Missing fields: {missing}"}

        prodi_id = data["prodi_id"]
        periode_id = data["periode_id"]
        jenis_dokumen = data["jenis_dokumen"]

        db = SessionLocal()
        try:
            prodi = db.execute(
                text("SELECT nama_prodi FROM m_prodi WHERE id = :prodi_id"),
                {"prodi_id": prodi_id},
            ).fetchone()

            periode = db.execute(
                text("SELECT nama_periode FROM m_periode_akademik WHERE id = :periode_id"),
                {"periode_id": periode_id},
            ).fetchone()

            nama_prodi = prodi.nama_prodi if prodi else f"Prodi #{prodi_id}"
            nama_periode = periode.nama_periode if periode else f"Periode #{periode_id}"

            indikator_data = db.execute(
                text("""
                    SELECT i.kode_indikator, i.nama_indikator, i.target, pi.nilai, pi.status
                    FROM trx_pemenuhan_indikator pi
                    JOIN m_indikator_akreditasi i ON i.id = pi.indikator_id
                    WHERE pi.prodi_id = :prodi_id AND pi.periode_id = :periode_id
                    ORDER BY i.kode_indikator
                """),
                {"prodi_id": prodi_id, "periode_id": periode_id},
            ).fetchall()

            # Generate narrative using Gemini LLM
            narasi = self._generate_narrative_with_gemini(nama_prodi, nama_periode, indikator_data, jenis_dokumen)

            sections = self._build_sections(nama_prodi, nama_periode, indikator_data, jenis_dokumen, narasi)

            docx_path = self._create_docx(nama_prodi, nama_periode, sections, jenis_dokumen)

            result = {
                "narasi": narasi,
                "sections": sections,
                "jenis_dokumen": jenis_dokumen,
                "prodi": nama_prodi,
                "periode": nama_periode,
                "file_path": docx_path,
            }

            db.execute(
                text("""
                    INSERT INTO agent_generator_history 
                        (prodi_id, periode_id, jenis_dokumen, judul, file_path, status, hasil_text, generated_by, created_at, updated_at)
                    VALUES (:prodi_id, :periode_id, :jenis, :judul, :file_path, :status, :hasil, :generated, NOW(), NOW())
                """),
                {
                    "prodi_id": prodi_id,
                    "periode_id": periode_id,
                    "jenis": jenis_dokumen,
                    "judul": f"{jenis_dokumen.upper()} - {nama_prodi} - {nama_periode}",
                    "file_path": docx_path,
                    "status": "selesai",
                    "hasil": json.dumps(result),
                    "generated": "agent",
                },
            )
            db.commit()

            self.log_execution(self.name, None, data, result)
            return result

        except Exception as e:
            logger.error(f"Generator error: {e}", exc_info=True)
            result = {"status": "error", "message": str(e)}
            self.log_execution(self.name, None, data, result, status="error", error_message=str(e))
            return result
        finally:
            db.close()

    def _generate_narrative_with_gemini(self, nama_prodi: str, nama_periode: str, indikator_data: list, jenis_dokumen: str) -> str:
        """
        Generate narrative using Gemini LLM based on indicator data
        """
        try:
            # Configure Gemini API
            genai.configure(api_key=GENAI_API_KEY)
            model = genai.GenerativeModel(GENAI_MODEL)
            
            # Prepare data for the prompt
            indikator_summary = []
            merah_count = 0
            kuning_count = 0
            hijau_count = 0
            
            for indikator in indikator_data:
                indikator_summary.append(
                    f"- {indikator.nama_indikator} (Kode: {indikator.kode_indikator}): "
                    f"Target {indikator.target}, Tercapai {indikator.nilai or '-'}, Status {indikator.status}"
                )
                if indikator.status == "merah":
                    merah_count += 1
                elif indikator.status == "kuning":
                    kuning_count += 1
                elif indikator.status == "hijau":
                    hijau_count += 1
            
            # Create prompt for Gemini
            prompt = f"""
            Buat narasi untuk dokumen {jenis_dokumen.upper()} Program Studi {nama_prodi} periode {nama_periode} berdasarkan data berikut:
            
            Jumlah Indikator: {len(indikator_data)}
            Status Capaian:
            - Merah (Perlu Tindakan Segera): {merah_count} indikator
            - Kuning (Perlu Ditingkatkan): {kuning_count} indikator
            - Hijau (Sudah Tercapai): {hijau_count} indikator
            
            Detail Indikator:
            {chr(10).join(indikator_summary)}
            
            Buat narasi yang profesional, kohesif, dan relevan dengan standar akreditasi. Narasi harus mencakup:
            1. Pendahuluan tentang pentingnya dokumen ini
            2. Analisis capaian indikator berdasarkan data di atas
            3. Identifikasi area yang perlu perbaikan
            4. Rekomendasi untuk peningkatan
            5. Penutup yang positif dan membangkitkan semangat
            
            Gunakan bahasa Indonesia yang formal tetapi mudah dipahami.
            """
            
            # Generate content using Gemini
            response = model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            logger.warning(f"Gemini generation failed, falling back to template: {e}")
            # Fallback to original template-based generation
            return self._generate_template_narrative(nama_prodi, nama_periode, indikator_data, jenis_dokumen)

    def _generate_template_narrative(self, nama_prodi: str, nama_periode: str, indikator_data: list, jenis_dokumen: str) -> str:
        """
        Original template-based narrative generation as fallback
        """
        # Build sections using original method to get the narrative
        sections = self._build_sections(nama_prodi, nama_periode, indikator_data, jenis_dokumen)
        return "\n\n".join(s["isi"] for s in sections)

    def _build_sections(self, prodi: str, periode: str, indikator_data: list, jenis: str, narasi: str = "") -> list:
        today = datetime.now().strftime("%d %B %Y")
        sections = []

        if jenis in ("led", "lkpt"):
            sections.append({
                "judul": "Cover",
                "isi": f"DOKUMEN {jenis.upper()}\nProgram Studi {prodi}\nPeriode {periode}\nDibuat: {today}",
            })

            # Use generated narrative if available, otherwise use template
            if narasi:
                sections.append({
                    "judul": "Pendahuluan dan Analisis",
                    "isi": narasi,
                })
            else:
                sections.append({
                    "judul": "Pendahuluan",
                    "isi": (
                        f"Dokumen {jenis.upper()} ini disusun untuk Program Studi {prodi} "
                        f"pada periode {periode}. Dokumen ini berisi ringkasan capaian indikator "
                        f"kinerja yang telah dicapai dalam periode tersebut."
                    ),
                })

            sections.append({
                "judul": "Capaian Indikator",
                "isi": self._format_indikator_table(indikator_data),
            })

            merah = [r for r in indikator_data if r.status == "merah"]
            kuning = [r for r in indikator_data if r.status == "kuning"]
            hijau = [r for r in indikator_data if r.status == "hijau"]

            analisa = []
            if merah:
                analisa.append(f"Terdapat {len(merah)} indikator berstatus merah yang memerlukan tindakan segera.")
            if kuning:
                analisa.append(f"Terdapat {len(kuning)} indikator berstatus kuning yang perlu ditingkatkan.")
            if hijau:
                analisa.append(f"Terdapat {len(hijau)} indikator berstatus hijau yang sudah tercapai.")
            sections.append({"judul": "Analisis", "isi": " ".join(analisa) if analisa else "Belum ada data capaian."})

            sections.append({
                "judul": "Penutup",
                "isi": f"Dokumen ini dihasilkan secara otomatis oleh sistem AI Akreditasi pada {today}.",
            })

        return sections

    def _format_indikator_table(self, rows: list) -> str:
        if not rows:
            return "Tidak ada data indikator."
        lines = ["| Kode | Indikator | Target | Tercapai | Status |", "|------|-----------|--------|----------|--------|"]
        for r in rows:
            target = r.target if r.target else "-"
            skor = r.nilai if r.nilai else "-"
            lines.append(f"| {r.kode_indikator} | {r.nama_indikator} | {target} | {skor} | {r.status} |")
        return "\n".join(lines)

    def _create_docx(self, prodi: str, periode: str, sections: list, jenis: str) -> str:
        doc = Document()
        
        title = doc.add_heading(f"DOKUMEN {jenis.upper()}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Program Studi: {prodi}")
        doc.add_paragraph(f"Periode: {periode}")
        doc.add_paragraph(f"Tanggal: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph()
        
        for section in sections:
            heading = doc.add_heading(section.get("judul", ""), level=1)
            content = section.get("isi", "")
            
            paragraphs = content.split("\n")
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para)
            
            doc.add_paragraph()
        
        from config import DOCS_OUTPUT_DIR
        output_dir = DOCS_OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"{jenis}_{prodi.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        logger.info(f"DOCX saved to: {filepath}")
        
        return filepath
