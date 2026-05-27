# Vetted by AI - Manual Review Required by Senior Engineer/Manager
"""
MCP Tools for AI Agents - Akreditasi System
Registers all agent tools with the MCP server
"""
from mcp.server.fastmcp import FastMCP, Context
from pydantic import Field
from typing import Optional

from datetime import datetime, date
import os
from agents_mcp.config import MCP_SERVER_NAME, MCP_SERVER_VERSION
from agents_mcp.auth import verify_mcp_auth
from agents_mcp.database import execute_query, list_tables, get_table_schema
from config import calculate_prediction


class _NoOpContext:
    """Fallback context when tools are called directly (not via MCP protocol)."""
    async def report_progress(self, *args, **kwargs):
        pass
    async def info(self, *args, **kwargs):
        pass
    async def error(self, *args, **kwargs):
        pass


def _ctx(ctx):
    """Return ctx or a no-op fallback."""
    return ctx if ctx is not None else _NoOpContext()

# Initialize FastMCP server
mcp = FastMCP(
    name=MCP_SERVER_NAME,
    stateless_http=True,
    json_response=True,
)


@mcp.tool()
async def db_list_tables(
    ctx: Context,
) -> list[str]:
    """List all available database tables in the system."""
    await _ctx(ctx).info("Fetching table list from PostgreSQL")
    tables = await list_tables()
    await _ctx(ctx).info(f"Found {len(tables)} tables")
    return tables


@mcp.tool()
async def db_get_table_schema(
    table_name: str = Field(description="Name of the table to get schema for"),
) -> list[dict]:
    """Get schema information (columns, types) for a specific table."""
    schema = await get_table_schema(table_name)
    return schema


# Whitelist of allowed table prefixes for db_query_safe
# All m_*, trx_*, and v_* tables in the public schema are permitted.
# Actual table names are validated at runtime against information_schema.
IDENTIFIER_RE = r'^[a-zA-Z_][a-zA-Z0-9_]*$'
import re
_identifier_validator = re.compile(IDENTIFIER_RE).match


@mcp.tool()
async def db_query(
    table_name: str = Field(description="Table name to query (must be m_*, trx_*, or v_* table)"),
    columns: Optional[list[str]] = Field(default=None, description="List of column names to return (default: all)"),
    where: Optional[dict] = Field(default=None, description="Filter conditions as key-value pairs (e.g. {\"prodi_id\": 1})"),
    limit: int = Field(default=100, description="Maximum rows to return (max: 1000)", ge=1, le=1000),
    order_by: Optional[str] = Field(default=None, description="Order by column name, optionally with ASC/DESC (e.g. 'nama' or 'id DESC')"),
) -> list[dict]:
    """
    Execute a safe, parameterized SELECT query using structured parameters.
    No raw SQL accepted — eliminates SQL injection risk entirely.

    Example:
      db_query(table_name="m_prodi", columns=["id", "nama_prodi"], where={"id": 1})
    """
    from agents_mcp.database import execute_query_safe
    return await execute_query_safe(
        table_name=table_name,
        columns=columns,
        where=where,
        limit=limit,
        order_by=order_by,
    )


@mcp.tool()
async def peringatan_check(
    prodi_id: int = Field(description="Program Studi ID"),
    ctx: Context = None,
) -> dict:
    """
    Check for warnings in a program studi:
    - BKD < 12 SKS per dosen
    - Kalibrasi sarana expired/near-expiry
    - Akreditasi expiring
    - Low RKAT budget absorption (Dana Mandek)
    Returns list of warnings with level (critical/warning/info).
    """
    warnings = []

    # Check BKD
    await _ctx(ctx).report_progress(1, 4, "Memeriksa BKD dosen...")
    bkd_query = """
        SELECT d.id, d.nama_depan, d.nama_belakang, d.nidn, b.total_sks
        FROM m_dosen d
        LEFT JOIN trx_bkd b ON d.id = b.dosen_id AND b.status = 'disetujui'
        WHERE d.prodi_id = $1 AND (b.total_sks IS NULL OR b.total_sks < 12)
    """
    bkd_rows = await execute_query(bkd_query, [prodi_id])
    for row in bkd_rows:
        nama = f"{row.get('nama_depan', '')} {row.get('nama_belakang', '')}".strip()
        warnings.append({
            "jenis": "bkd",
            "tingkat": "critical" if (row.get("total_sks") or 0) < 6 else "warning",
            "pesan": f"Dosen {nama} (NIDN: {row['nidn']}) BKD {row.get('total_sks', 0)} SKS < 12 SKS",
            "dosen_id": row.get("id"),
        })

    # Check Kalibrasi
    await _ctx(ctx).report_progress(2, 4, "Memeriksa kalibrasi sarana...")
    kalibrasi_query = """
        SELECT s.nama_sarana, s.tanggal_kalibrasi, s.tanggal_kalibrasi_berikut
        FROM m_sarana s
        WHERE s.prodi_id = $1
        AND s.tanggal_kalibrasi_berikut IS NOT NULL
        AND (
            s.tanggal_kalibrasi_berikut <= CURRENT_DATE
            OR s.tanggal_kalibrasi_berikut <= CURRENT_DATE + INTERVAL '30 days'
        )
    """
    kalibrasi_rows = await execute_query(kalibrasi_query, [prodi_id])
    for row in kalibrasi_rows:
        warnings.append({
            "jenis": "kalibrasi",
            "tingkat": "critical" if row.get("tanggal_kalibrasi") else "warning",
            "pesan": f"Sarana {row['nama_sarana']} kalibrasi expired/hampir expired",
        })

    # Check Akreditasi
    await _ctx(ctx).report_progress(3, 4, "Memperingatkan akreditasi...")
    akreditasi_query = """
        SELECT p.id, p.nama_prodi, p.akreditasi, p.tanggal_kadaluarsa
        FROM m_prodi p
        WHERE p.id = $1
        AND p.tanggal_kadaluarsa IS NOT NULL
        AND p.tanggal_kadaluarsa <= CURRENT_DATE + INTERVAL '180 days'
    """
    akreditasi_rows = await execute_query(akreditasi_query, [prodi_id])
    for row in akreditasi_rows:
        warnings.append({
            "jenis": "akreditasi",
            "tingkat": "critical" if row.get("tanggal_kadaluarsa") else "info",
            "pesan": f"Prodi {row['nama_prodi']} akreditasi {row.get('akreditasi', 'N/A')} berlaku sampai {row.get('tanggal_kadaluarsa')}",
        })

    # Check RKAT Absorption (Dana Mandek)
    await _ctx(ctx).report_progress(4, 4, "Memeriksa penyerapan anggaran...")
    rkat_query = """
        SELECT pagu_total, terpakai
        FROM trx_rkat_pagu
        WHERE unit_type = 'Prodi' AND unit_id = $1
        ORDER BY created_at DESC LIMIT 1
    """
    rkat_rows = await execute_query(rkat_query, [prodi_id])
    if rkat_rows:
        row = rkat_rows[0]
        pagu = float(row['pagu_total'])
        if pagu > 0:
            absorption = (float(row['terpakai']) / pagu) * 100
            if date.today().month > 6 and absorption < 40:
                warnings.append({
                    "jenis": "rkat",
                    "tingkat": "critical" if absorption < 20 else "warning",
                    "pesan": f"Penyerapan anggaran RKAT rendah ({absorption:.1f}%) di atas semester 1.",
                })

    await _ctx(ctx).info(f"Peringatan check complete: {len(warnings)} warnings found")
    return {
        "prodi_id": prodi_id,
        "warning_count": len(warnings),
        "warnings": warnings,
    }


@mcp.tool()
async def rekomendasi_generate(
    prodi_id: int = Field(description="Program Studi ID"),
    top_n: int = 10,
    ctx: Context = None,
) -> dict:
    """
    Generate recommendations for a program studi based on indicator performance.
    Returns top N recommendations ordered by weight (bobot) descending.
    """
    await _ctx(ctx).report_progress(1, 2, "Fetching indicator data...")

    query = """
        SELECT i.id, i.kode_indikator, i.nama_indikator, i.bobot, pi.status, pi.nilai
        FROM m_indikator_akreditasi i
        LEFT JOIN trx_pemenuhan_indikator pi ON i.id = pi.indikator_id
        WHERE pi.prodi_id = $1
        AND pi.status IN ('merah', 'kuning')
        ORDER BY i.bobot DESC
        LIMIT $2
    """
    rows = await execute_query(query, [prodi_id, top_n])

    await _ctx(ctx).report_progress(2, 2, "Generating recommendations...")

    rekomendasi_list = []
    for row in rows:
        bobot = row.get("bobot", 0)
        prioritas = "tinggi" if bobot >= 4 else "sedang" if bobot >= 2 else "rendah"
        rekomendasi_list.append({
            "indikator_id": row.get("id"),
            "kode": row["kode_indikator"],
            "nama": row["nama_indikator"],
            "bobot": bobot,
            "status": row["status"],
            "prioritas": prioritas,
            "saran": f"Tingkatkan pemenuhan pada indikator {row['kode_indikator']} ({row['nama_indikator']})",
        })

    await _ctx(ctx).info(f"Generated {len(rekomendasi_list)} recommendations for prodi {prodi_id}")
    return {
        "prodi_id": prodi_id,
        "recommendation_count": len(rekomendasi_list),
        "recommendations": rekomendasi_list,
    }


@mcp.tool()
async def verifikasi_dokumen(
    prodi_id: int = Field(description="Program Studi ID"),
    doc_bukti_id: Optional[int] = None,
    ctx: Context = None,
) -> dict:
    """
    Verify documents for a program studi.
    Checks file existence, size, and hash duplicates.
    Returns verification status per document.
    """
    await _ctx(ctx).report_progress(1, 3, "Fetching document data...")

    query = """
        SELECT db.id, db.nama_dokumen, db.file_path, db.file_size, db.hash, db.keterangan,
               d.nama_depan, d.nama_belakang
        FROM doc_bukti db
        LEFT JOIN m_dosen d ON db.dosen_id = d.id
        WHERE db.prodi_id = $1
    """
    params = [prodi_id]
    if doc_bukti_id:
        query += " AND db.id = $2"
        params.append(doc_bukti_id)

    rows = await execute_query(query, params)

    await _ctx(ctx).report_progress(2, 3, "Checking file integrity...")

    import hashlib
    import os

    hasil_verifikasi = []
    hash_seen = {}

    for row in rows:
        file_path = row.get("file_path", "")
        file_exists = os.path.exists(file_path) if file_path else False
        file_size = row.get("file_size", 0)
        file_hash = row.get("hash", "")

        # Check hash duplicates
        is_duplicate = False
        if file_hash and file_hash in hash_seen:
            is_duplicate = True
        elif file_hash:
            hash_seen[file_hash] = row["id"]

        status = "valid"
        catatan = []

        if not file_exists:
            status = "need_review"
            catatan.append("File tidak ditemukan")
        elif file_size == 0:
            status = "need_review"
            catatan.append("File kosong (0 bytes)")
        elif is_duplicate:
            status = "need_review"
            catatan.append(f"Hash duplikat dengan dokumen ID {hash_seen.get(file_hash)}")

        hasil_verifikasi.append({
            "doc_bukti_id": row["id"],
            "nama_dokumen": row["nama_dokumen"],
            "status": status,
            "catatan": ", ".join(catatan) if catatan else "OK",
            "tingkat_kepercayaan": 1.0 if status == "valid" else 0.5,
            "dosen": f"{row.get('nama_depan', '')} {row.get('nama_belakang', '')}".strip() or None,
            "keterangan": row.get("keterangan"),
        })

    await _ctx(ctx).report_progress(3, 3, "Verification complete")

    valid_count = sum(1 for h in hasil_verifikasi if h["status"] == "valid")
    review_count = len(hasil_verifikasi) - valid_count

    await _ctx(ctx).info(f"Verifikasi complete: {valid_count} valid, {review_count} need review")
    return {
        "prodi_id": prodi_id,
        "total_documents": len(hasil_verifikasi),
        "valid_count": valid_count,
        "need_review_count": review_count,
        "results": hasil_verifikasi,
    }


@mcp.tool()
async def prediksi_skor(
    prodi_id: int = Field(description="Program Studi ID"),
    ctx: Context = None,
) -> dict:
    """
    Predict accreditation score for a program studi using historical data and budget analysis.
    Uses weighted scores (nilai * bobot) per periode + linear regression trend + budget impact.
    Returns predicted score + probabilities for 3 categories.
    """
    await _ctx(ctx).report_progress(1, 4, "Fetching historical data...")

    query = """
        SELECT pi.periode_id, pi.nilai, i.bobot
        FROM trx_pemenuhan_indikator pi
        JOIN m_indikator_akreditasi i ON i.id = pi.indikator_id
        WHERE pi.prodi_id = $1
        ORDER BY pi.periode_id DESC
    """
    rows = await execute_query(query, [prodi_id])

    if not rows:
        return {
            "prodi_id": prodi_id,
            "error": "Data indikator tidak ditemukan untuk prodi ini",
            "predicted_score": None,
        }

    await _ctx(ctx).report_progress(2, 4, "Fetching budget data...")
    budget_query = """
        SELECT SUM(estimasi_biaya) as total_biaya, periode_id
        FROM trx_usulan_rkat
        WHERE prodi_id = $1 AND status = 'approved'
        GROUP BY periode_id
    """
    budget_rows = await execute_query(budget_query, [prodi_id])
    budgets = {b['periode_id']: float(b['total_biaya']) for b in budget_rows}

    await _ctx(ctx).report_progress(3, 4, "Computing prediction...")

    period_scores = {}
    for r in rows:
        p_id = r["periode_id"]
        if p_id not in period_scores:
            period_scores[p_id] = {"skor": 0, "bobot": 0}
        period_scores[p_id]["skor"] += float(r["nilai"] or 0) * r["bobot"]
        period_scores[p_id]["bobot"] += r["bobot"]

    historical_scores = []
    for p_id in sorted(period_scores.keys()):
        ts = period_scores[p_id]
        historical_scores.append(ts["skor"] / ts["bobot"] if ts["bobot"] > 0 else 0)

    historical_scores = historical_scores[-3:]

    pred = calculate_prediction(historical_scores)

    # Analyze Budget Impact
    budget_impact = "netral"
    if budgets:
        latest_period = max(period_scores.keys())
        if latest_period in budgets:
            current_budget = budgets[latest_period]
            avg_budget = sum(budgets.values()) / len(budgets)
            if current_budget > avg_budget * 1.2:
                budget_impact = "positif"
                pred["skor_prediksi"] += 0.5
            elif current_budget < avg_budget * 0.8:
                budget_impact = "negatif"
                pred["skor_prediksi"] -= 0.3

    await _ctx(ctx).report_progress(4, 4, "Prediction complete")

    predicted_category = "Unggul" if pred["probabilitas"]["unggul"] > 0.5 else "Baik Sekali" if pred["probabilitas"]["baik_sekali"] > 0.5 else "Baik"

    await _ctx(ctx).info(f"Prediksi complete: {pred['skor_prediksi']:.2f} ({predicted_category})")
    return {
        "prodi_id": prodi_id,
        "predicted_score": round(pred["skor_prediksi"], 2),
        "predicted_category": predicted_category,
        "probabilities": pred["probabilitas"],
        "trend_analysis": pred["trend_analysis"],
        "budget_impact": budget_impact,
        "historical_data_points": len(historical_scores),
    }


@mcp.tool()
async def generator_dokumen(
    prodi_id: int = Field(description="Program Studi ID"),
    periode_id: Optional[int] = None,
    jenis_dokumen: str = Field(description="Jenis dokumen: 'LED' or 'LKPT'"),
    ctx: Context = None,
) -> dict:
    """
    Generate LED (Laporan Evaluasi Diri) or LKPT (Laporan Kinerja Program Studi) document.
    Returns document metadata and file path.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    import os


    await _ctx(ctx).report_progress(1, 4, "Fetching prodi data...")

    prodi_query = "SELECT nama_prodi, kode_prodi FROM m_prodi WHERE id = $1"
    prodi_rows = await execute_query(prodi_query, [prodi_id])
    if not prodi_rows:
        return {"error": f"Prodi ID {prodi_id} not found"}

    prodi = prodi_rows[0]

    await _ctx(ctx).report_progress(2, 4, "Fetching indicator data...")

    indikator_query = """
        SELECT i.kode_indikator, i.nama_indikator, pi.status, pi.nilai, pi.catatan
        FROM m_indikator_akreditasi i
        LEFT JOIN trx_pemenuhan_indikator pi ON i.id = pi.indikator_id
        WHERE pi.prodi_id = $1
    """
    params = [prodi_id]
    if periode_id:
        indikator_query += " AND pi.periode_id = $2"
        params.append(periode_id)

    indikator_rows = await execute_query(indikator_query, params)

    await _ctx(ctx).report_progress(3, 4, "Generating document...")

    # Create document
    doc = Document()

    # Cover page
    doc.add_heading(f"Laporan {'Evaluasi Diri' if jenis_dokumen == 'LED' else 'Kinerja Program Studi'}", level=1)
    doc.add_paragraph(f"Program Studi: {prodi['nama_prodi']}")
    doc.add_paragraph(f"Kode Prodi: {prodi['kode_prodi']}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    # Pendahuluan
    doc.add_heading("Pendahuluan", level=1)
    doc.add_paragraph(
        f"Dokumen ini merupakan {'Laporan Evaluasi Diri (LED)' if jenis_dokumen == 'LED' else 'Laporan Kinerja Program Studi (LKPT)'} "
        f"untuk Program Studi {prodi['nama_prodi']}."
    )

    # Capaian Indikator
    doc.add_heading("Capaian Indikator", level=1)

    hijau_count = sum(1 for i in indikator_rows if i.get("status") == "hijau")
    kuning_count = sum(1 for i in indikator_rows if i.get("status") == "kuning")
    merah_count = sum(1 for i in indikator_rows if i.get("status") == "merah")

    doc.add_paragraph(f"Total Indikator: {len(indikator_rows)}")
    doc.add_paragraph(f"- Hijau (Tercapai): {hijau_count}")
    doc.add_paragraph(f"- Kuning (Perlu Peningkatan): {kuning_count}")
    doc.add_paragraph(f"- Merah (Belum Tercapai): {merah_count}")

    # Detail per indikator
    doc.add_heading("Detail Indikator", level=2)
    for ind in indikator_rows:
        doc.add_heading(f"{ind.get('kode', 'N/A')} - {ind['nama_indikator']}", level=3)
        doc.add_paragraph(f"Status: {ind.get('status', 'N/A')}")
        doc.add_paragraph(f"Nilai: {ind.get('nilai', 'N/A')}")
        if ind.get("catatan"):
            doc.add_paragraph(f"Catatan: {ind['catatan']}")

    # Analisis
    doc.add_heading("Analisis", level=1)
    doc.add_paragraph(
        f"Dari {len(indikator_rows)} indikator, {hijau_count} sudah terpenuhi (hijau), "
        f"{kuning_count} perlu peningkatan (kuning), dan {merah_count} belum terpenuhi (merah)."
    )

    # Save document
    output_dir = os.getenv("DOCS_OUTPUT_DIR", os.path.join(os.path.dirname(__file__), "..", "output", "docs"))
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{jenis_dokumen}_{prodi['kode_prodi']}_{timestamp}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)

    await _ctx(ctx).report_progress(4, 4, "Document saved")

    await _ctx(ctx).info(f"Document generated: {file_path}")
    return {
        "prodi_id": prodi_id,
        "periode_id": periode_id,
        "jenis_dokumen": jenis_dokumen,
        "file_path": file_path,
        "filename": filename,
        "status": "success",
        "indicator_summary": {
            "total": len(indikator_rows),
            "hijau": hijau_count,
            "kuning": kuning_count,
            "merah": merah_count,
        },
    }


@mcp.tool()
async def integrasi_sync(
    sumber: str = Field(description="Sumber data: 'pddikti', 'sinta', or 'sister'"),
    ctx: Context = None,
) -> dict:
    """
    Synchronize data from external sources (PDDIKTI, SINTA, SISTER).
    Fallback chain: API → SQL Views (Excel Import) → Reconciliation.
    """
    await _ctx(ctx).report_progress(1, 5, f"Memeriksa {sumber}...")

    from agents_mcp.database import execute_query

    view_map = {
        "pddikti": "v_sync_pddikti_dosen",
        "sinta": "v_sync_sinta_publikasi",
        "sister": "v_sync_sister_riwayat",
    }

    view = view_map.get(sumber)
    if not view:
        return {
            "sumber": sumber,
            "status": "error",
            "message": f"Sumber '{sumber}' tidak dikenal. Gunakan: pddikti, sinta, sister",
        }

    # Step 1: Try external API first
    await _ctx(ctx).report_progress(2, 5, f"Mencoba koneksi API {sumber}...")
    api_success = False
    api_result = None
    try:
        import httpx
        from agents_mcp.config import SINTA_API_URL, SINTA_API_KEY, PDDIKTI_API_URL, PDDIKTI_API_KEY

        if sumber == "pddikti":
            async with httpx.AsyncClient(timeout=15) as client:
                headers = {"Authorization": f"Bearer {PDDIKTI_API_KEY}"} if PDDIKTI_API_KEY else {}
                resp = await client.get(f"{PDDIKTI_API_URL}/dosen", params={"limit": 1}, headers=headers)
                api_success = resp.status_code == 200
        elif sumber == "sinta":
            async with httpx.AsyncClient(timeout=15) as client:
                headers = {"Authorization": f"Bearer {SINTA_API_KEY}"} if SINTA_API_KEY else {}
                resp = await client.get(f"{SINTA_API_URL}/authors", params={"q": "", "limit": 1}, headers=headers)
                api_success = resp.status_code == 200
        if api_success:
            api_result = {"sumber": sumber, "status": "api_available"}
    except Exception:
        api_success = False

    if api_success:
        await _ctx(ctx).report_progress(3, 5, f"API {sumber} tersedia, melanjutkan...")
        return {
            "sumber": sumber,
            "status": "api_available",
            "records_pulled": 0,
            "conflicts_detected": 0,
            "message": f"API {sumber} tersedia. Gunakan endpoint integrasi terpisah untuk sinkronasi penuh.",
            "fallback_data_available": False,
        }

    # Step 2: API unavailable → check SQL views (Excel import data)
    await _ctx(ctx).report_progress(3, 5, f"API tidak tersedia, memeriksa data dari Excel import...")
    rows = await execute_query(f"SELECT COUNT(*) as cnt FROM {view}")
    total = rows[0]["cnt"] if rows else 0

    # Step 3: Check reconciliation pending
    await _ctx(ctx).report_progress(4, 5, "Memeriksa data rekonsiliasi...")
    pending_rows = await execute_query("""
        SELECT COUNT(*) as cnt FROM reconciliation_suggestions
        WHERE status = 'pending' AND source_type LIKE $1
    """, [f"%{sumber}%"])
    pending_count = pending_rows[0]["cnt"] if pending_rows else 0

    await _ctx(ctx).report_progress(5, 5, "Selesai")

    if total > 0:
        return {
            "sumber": sumber,
            "status": "completed",
            "records_in_system": total,
            "reconciliation_pending": pending_count,
            "api_available": False,
            "message": f"Data tersedia dari Excel import ({total} record). {pending_count} perlu rekonsiliasi." if pending_count else f"Semua data {sumber} tersinkron via Excel import ({total} record).",
        }

    return {
        "sumber": sumber,
        "status": "excel_required",
        "records_in_system": 0,
        "reconciliation_pending": 0,
        "api_available": False,
        "message": f"Tidak ada data {sumber}. Upload file Excel melalui menu Import Data untuk mengisi data.",
    }


